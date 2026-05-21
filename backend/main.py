from fastapi import FastAPI, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from lxml import etree
from latex2mathml.converter import convert as latex_to_mathml

import re
import os
import uuid
import logging

# =========================================
# LOGGING
# =========================================

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)

logger = logging.getLogger(__name__)

# =========================================
# FASTAPI
# =========================================

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
allow_origins=[
    "https://www.mspkchatgpttoword.com",
    "https://mspkchatgpttoword.com"
],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# =========================================
# INPUT MODEL
# =========================================

class TextInput(BaseModel):
    text: str = Field(..., min_length=1, max_length=500000)
    title: str = Field(
        default="CONVERTED INTO WORD"
    )

# =========================================
# PAGE BORDER
# =========================================

def add_page_border(doc):

    for section in doc.sections:

        sectPr = section._sectPr

        pgBorders = OxmlElement('w:pgBorders')

        pgBorders.set(
            qn('w:offsetFrom'),
            'page'
        )

        for border_name in [
            'top',
            'left',
            'bottom',
            'right'
        ]:

            border_el = OxmlElement(
                f'w:{border_name}'
            )

            border_el.set(
                qn('w:val'),
                'single'
            )

            border_el.set(
                qn('w:sz'),
                '4'
            )

            border_el.set(
                qn('w:space'),
                '24'
            )

            border_el.set(
                qn('w:color'),
                '0a2540'
            )

            pgBorders.append(border_el)

        sectPr.append(pgBorders)

# =========================================
# REMOVE AI GARBAGE
# =========================================

def remove_ai_garbage(text):

    garbage_patterns = [

        r"^✅\s+",
        r"^\{\s*$",
        r'^\}\s*$',
        r'^"text"\s*:',
        r"^Convert this into",
        r"^Create\s+",
        r"^Design\s+",
        r"if you want next",
        r"run.*pip.*install",

    ]

    cleaned_lines = []

    for line in text.splitlines():

        skip = False

        for pattern in garbage_patterns:

            if re.search(
                pattern,
                line,
                re.IGNORECASE
            ):

                skip = True
                break

        if not skip and line.strip():

            cleaned_lines.append(line)

    return "\n".join(cleaned_lines)

# =========================================
# CLEAN TEXT
# =========================================

def clean_text(text):

    text = re.sub(
        r"^#+\s*",
        "",
        text
    )

    text = text.replace(
        "**",
        ""
    )

    return text.strip()

# =========================================
# FIX TEXT FRACTIONS
# =========================================

def fix_text_fractions(eq):

    pattern = r'\\frac\s*\{([^{}]+)\}\s*\{([^{}]+)\}'

    matches = re.findall(pattern, eq)

    for num, den in matches:

        new_num = num
        new_den = den

        if (
            " " in num
            and
            r"\text" not in num
            and
            "\\" not in num
        ):

            new_num = rf"\text{{{num}}}"

        if (
            " " in den
            and
            r"\text" not in den
            and
            "\\" not in den
        ):

            new_den = rf"\text{{{den}}}"

        old = rf"\frac{{{num}}}{{{den}}}"

        new = rf"\frac{{{new_num}}}{{{new_den}}}"

        eq = eq.replace(old, new)

    return eq

# =========================================
# CLEAN EQUATION
# =========================================

def simplify_nested_fractions(eq):

    # DOUBLE NESTED FRACTION

    pattern = (
        r'\\frac\s*\{'
        r'\s*\\frac\{([^{}]+)\}\{([^{}]+)\}'
        r'\s*\}'
        r'\s*\{'
        r'\s*\\frac\{([^{}]+)\}\{([^{}]+)\}'
        r'\s*\}'
    )

    matches = re.findall(pattern, eq)

    for a, b, c, d in matches:

        # CONVERT:
        # (a/b)/(c/d)
        # INTO:
        # (a*d)/(b*c)

        replacement = (
            rf'\\frac{{({a})({d})}}{{({b})({c})}}'
        )

        old = (
            rf'\frac{{\frac{{{a}}}{{{b}}}}}'
            rf'{{\frac{{{c}}}{{{d}}}}}'
        )

        eq = eq.replace(old, replacement)

    return eq

def clean_equation(eq):

    eq = eq.strip()

    # REMOVE WRAPPERS

    eq = eq.replace("\\[", "")
    eq = eq.replace("\\]", "")

    eq = eq.replace("\\(", "")
    eq = eq.replace("\\)", "")

    eq = eq.replace("$$", "")

    # ALIGN SUPPORT

    if (
        r"\begin{aligned}" in eq
        or
        r"\begin{align}" in eq
    ):

        eq = eq.replace(
            r"\begin{aligned}",
            ""
        )

        eq = eq.replace(
            r"\end{aligned}",
            ""
        )

        eq = eq.replace(
            r"\begin{align}",
            ""
        )

        eq = eq.replace(
            r"\end{align}",
            ""
        )

        eq = eq.replace("&", "")

        eq = eq.replace(
            r"\\",
            " "
        )

    # KEEP LATEX SAFE

    eq = eq.replace("\\\\", "\\")

    # FIX TEXT FRACTIONS

    eq = fix_text_fractions(eq)

    # SYMBOL FIXES

    eq = eq.replace(
        "×",
        r"\times "
    )

    eq = eq.replace("₹", "")

    # REMOVE UNSUPPORTED

    eq = eq.replace("{text}", "")

    eq = eq.replace(r"\boxed", "")

    eq = eq.replace(r"\left", "")
    eq = eq.replace(r"\right", "")

    eq = eq.replace(r"\Big", "")
    eq = eq.replace(r"\big", "")

    eq = eq.replace(r"\displaystyle", "")

    # NORMALIZE

    eq = re.sub(
        r'\n+',
        ' ',
        eq
    )

    eq = re.sub(
        r'\s+',
        ' ',
        eq
    )
    # SIMPLIFY DEEP NESTED FRACTIONS

    eq = simplify_nested_fractions(eq)

    return eq.strip()

# =========================================
# EQUATION DETECTION
# =========================================

def is_equation_line(line):

    patterns = [

        r"\\frac",
        r"\\sum",
        r"\\sqrt",
        r"\\times",
        r"\\begin",
        r"\\end",
        r"\\int",
        r"\\prod",
        r"\\lim",
        r"\\partial",

        r"Ke\s*=",
        r"Kp\s*=",
        r"Kd\s*=",
        r"Re\s*=",
        r"WACC\s*=",
        r"NPV\s*=",
        r"IRR\s*=",
        r"ROE\s*=",

        r"_\{?.+?\}?",

    ]

    for pattern in patterns:

        try:

            if re.search(
                pattern,
                line,
                re.IGNORECASE
            ):
                return True

        except re.error:
            continue

    return False

# =========================================
# HEADING DETECTION
# =========================================

def is_heading(line):

    return (

        (
            line.strip().isupper()
            and
            len(line.strip()) > 3
        )

        or

        (
            line.strip().endswith(":")
            and
            len(line.strip()) < 70
        )

        or

        re.match(
            r"^(WN|worked note)\s+\d+",
            line,
            re.IGNORECASE
        )

        or

        re.match(
            r"^#+\s+",
            line
        )
    )

# =========================================
# TABLE DETECTION
# =========================================

def is_table_line(line):

    return line.count("|") >= 2

# =========================================
# WORD EQUATION ENGINE
# =========================================

def add_word_equation(
    paragraph,
    equation
):

    cleaned_eq = clean_equation(
        equation
    )

    # ESCAPE % ONLY FOR REAL EQUATIONS

    cleaned_eq = re.sub(
        r'(?<!\\)%',
        r'\\%',
        cleaned_eq
    )

    # DETECT VERY DEEP NESTED FRACTIONS

    frac_depth = cleaned_eq.count(r"\\frac")

    if frac_depth >= 3:

        fallback_eq = cleaned_eq

        fallback_eq = fallback_eq.replace(
            r"\\frac",
            "/"
        )

        fallback_eq = fallback_eq.replace(
            "{",
            "("
        )

        fallback_eq = fallback_eq.replace(
            "}",
            ")"
        )

        fallback_eq = fallback_eq.replace(
            r"\\times",
            "×"
        )

        fallback_eq = fallback_eq.replace(
            r"\\sqrt",
            "√"
        )

        fallback_eq = re.sub(
            r'\\\\[a-zA-Z]+',
            '',
            fallback_eq
        )

        run = paragraph.add_run(
            fallback_eq
        )

        run.font.name = "Times New Roman"

        run.italic = True

        return
    try:

        mathml = latex_to_mathml(
            cleaned_eq
        )

        if 'xmlns=' not in mathml:

            mathml = mathml.replace(

                "<math>",

                '<math xmlns="http://www.w3.org/1998/Math/MathML">'

            )

        xslt_path = os.path.join(
            os.getcwd(),
            "MML2OMML.XSL"
        )

        if not os.path.exists(xslt_path):

            run = paragraph.add_run(
                cleaned_eq
            )

            run.italic = True

            run.font.name = "Courier New"

            return

        xslt = etree.parse(xslt_path)

        transform = etree.XSLT(xslt)

        mathml_dom = etree.fromstring(
            mathml.encode("utf-8")
        )

        omml_dom = transform(mathml_dom)

        paragraph._element.append(
            omml_dom.getroot()
        )

    except Exception as e:

        logger.error(
            f"Equation Error: {e}"
        )

        run = paragraph.add_run(
            cleaned_eq
        )

        run.italic = True

        run.font.name = "Courier New"

# =========================================
# TABLE ENGINE
# =========================================

def add_table(doc, lines):

    rows = []

    for line in lines:

        if re.search(
            r"[-:]{3,}",
            line
        ):
            continue

        raw_cols = line.split("|")

        # PRESERVE EMPTY CELLS

        if (
            raw_cols
            and
            raw_cols[0].strip() == ""
        ):
            raw_cols = raw_cols[1:]

        if (
            raw_cols
            and
            raw_cols[-1].strip() == ""
        ):
            raw_cols = raw_cols[:-1]

        cols = [

            clean_text(c.strip())

            for c in raw_cols

        ]

        if cols:
            rows.append(cols)

    if not rows:
        return

    max_cols = max(
        len(r)
        for r in rows
    )

    table = doc.add_table(

        rows=len(rows),

        cols=max_cols

    )

    table.style = "Table Grid"

    table.autofit = True

    for i, row in enumerate(rows):

        while len(row) < max_cols:
            row.append("")

        for j, cell_text in enumerate(row):

            cell_obj = table.cell(i, j)

            para = cell_obj.paragraphs[0]

            para.paragraph_format.space_before = Pt(0)

            para.paragraph_format.space_after = Pt(2)

            para.paragraph_format.line_spacing = 1.0

            cleaned_cell = cell_text.strip()

            # REMOVE OUTER [ ] ONLY

            if (
                cleaned_cell.startswith("[")
                and
                cleaned_cell.endswith("]")
            ):

                cleaned_cell = cleaned_cell[
                    1:-1
                ].strip()

            # REMOVE ESCAPED SPACES
            # ONLY FOR NORMAL TEXT EQUATIONS

            if (
                r"\frac" not in cleaned_cell
            ):

                cleaned_cell = cleaned_cell.replace(
                    r"\ ",
                    " "
                )            

            # TRUE LATEX EQUATION ONLY

            latex_equation = (

                r"\frac" in cleaned_cell

                or

                r"\sqrt" in cleaned_cell

                or

                r"\sum" in cleaned_cell

                or

                r"\int" in cleaned_cell

                or

                r"\beta" in cleaned_cell

                or

                r"\alpha" in cleaned_cell

            )

            # EQUATION CELL

            if latex_equation:

                para.alignment = (
                    WD_PARAGRAPH_ALIGNMENT.CENTER
                )

                try:

                    add_word_equation(
                        para,
                        cleaned_cell
                    )

                except Exception as e:

                    logger.error(
                        f"Table Equation Error: {e}"
                    )

                    run = para.add_run(
                        cleaned_cell
                    )

                    run.font.name = (
                        "Courier New"
                    )

                    run.font.size = Pt(9)

            # NORMAL CELL

            else:

                normal_text = cleaned_cell

                # REMOVE [ ]

                if (
                    normal_text.startswith("[")
                    and
                    normal_text.endswith("]")
                ):

                    normal_text = normal_text[
                        1:-1
                    ].strip()

                # REMOVE ESCAPED SPACES

                normal_text = normal_text.replace(
                    r"\ ",
                    " "
                )

                if re.search(
                    r'[\d%₹,.]+',
                    normal_text
                ):

                    para.alignment = (
                        WD_PARAGRAPH_ALIGNMENT.CENTER
                    )

                else:

                    para.alignment = (
                        WD_PARAGRAPH_ALIGNMENT.LEFT
                    )

                match = re.search(
                    r'([A-Za-z]+)\(?_\{([^{}]+)\}\)?',
                    normal_text
                )         
                if match:

                    main_text = match.group(1)

                    sub_text = match.group(2)

                    run_main = para.add_run(
                        main_text
                    )

                    run_main.font.name = (
                        "Times New Roman"
                    )

                    run_main.font.size = Pt(12)

                    run_sub = para.add_run(
                        sub_text
                    )

                    run_sub.font.subscript = True

                    run_sub.font.name = (
                        "Times New Roman"
                    )

                    run_sub.font.size = Pt(10)

                else:

                    run = para.add_run(
                        normal_text
                    )

                    run.font.name = (
                        "Times New Roman"
                    )

                    run.font.size = Pt(12)
                
                # HEADER ROW

                if i == 0:

                    run.bold = True

                    run.font.color.rgb = (
                        RGBColor(
                            255,
                            255,
                            255
                        )
                    )

                    tcPr = (
                        cell_obj._tc
                        .get_or_add_tcPr()
                    )

                    shd = OxmlElement('w:shd')

                    shd.set(
                        qn('w:fill'),
                        '0A2540'
                    )

                    tcPr.append(shd)

    doc.add_paragraph()

# =========================================
# MAIN ROUTE
# =========================================

@app.post("/generate-docx")

async def generate_docx(
    data: TextInput,
    background_tasks: BackgroundTasks
):

    text = remove_ai_garbage(
        data.text
    )

    doc = Document()

    section = doc.sections[0]

    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    add_page_border(doc)

    style = doc.styles["Normal"]

    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # TITLE

    title = doc.add_paragraph()

    title.alignment = (
        WD_PARAGRAPH_ALIGNMENT.CENTER
    )

    title_run = title.add_run(
        data.title
    )

    title_run.bold = True

    title_run.font.size = Pt(16)

    title_run.font.color.rgb = (
        RGBColor(
            10,
            37,
            64
        )
    )

    title.paragraph_format.space_after = Pt(12)

    # PROCESS CONTENT

    lines = text.split("\n")

    i = 0

    while i < len(lines):

        line = lines[i].strip()

        # EMPTY

        if not line:

            doc.add_paragraph()

            i += 1

            continue

        # TABLE

        if is_table_line(line):

            table_lines = []

            while (

                i < len(lines)

                and

                is_table_line(lines[i])

            ):

                table_lines.append(
                    lines[i]
                )

                i += 1

            add_table(
                doc,
                table_lines
            )

            continue

        # =====================================
        # MULTILINE EQUATION
        # =====================================

        clean_line = line.strip()

        if (
            clean_line in ["[", r"\["]
            or
            clean_line.endswith("[")
            or
            clean_line.endswith(r"\[")
        ):

            full_eq = ""

            i += 1

            while i < len(lines):

                current = lines[i].strip()

                # SKIP EMPTY LINES

                if not current.strip():

                    i += 1
                    continue

                # END EQUATION

                if (
                    current.strip() == "]"
                    or
                    current.strip() == r"\]"
                ):
                    break

                full_eq += "\n" + current

                i += 1

            p = doc.add_paragraph()

            p.alignment = (
                WD_PARAGRAPH_ALIGNMENT.CENTER
            )

            add_word_equation(
                p,
                full_eq
            )

            i += 1

            continue

          # PURE EQUATION LINE ONLY

        if (

            is_equation_line(line)

            and

            len(line.split()) <= 12

            and

            "$" not in line

        ):

            p = doc.add_paragraph()

            p.alignment = (
                WD_PARAGRAPH_ALIGNMENT.CENTER
            )

            add_word_equation(
                p,
                line
            )

            i += 1

            continue

        # PROFESSIONAL DIVIDER

        if line.strip() == "---":

            p = doc.add_paragraph()

            run = p.add_run(
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━"
            )
            
            run.bold = True

            run.font.color.rgb = RGBColor(
                120,
                120,
                120
            )

            run.font.size = Pt(16)

            p.alignment = (
                WD_PARAGRAPH_ALIGNMENT.CENTER
            )

            p.paragraph_format.space_before = Pt(8)

            p.paragraph_format.space_after = Pt(8)

            i += 1

            continue

        # HEADING

        line_clean = clean_text(line)

        if is_heading(line_clean):

            p = doc.add_paragraph()

            run = p.add_run(
                line_clean
            )

            run.bold = True

            run.font.size = Pt(13)

            run.font.color.rgb = (
                RGBColor(
                    10,
                    37,
                    64
                )
            )

            p.paragraph_format.space_before = Pt(8)

            p.paragraph_format.space_after = Pt(4)

            i += 1

            continue

        # NORMAL TEXT

        p = doc.add_paragraph()

        p.paragraph_format.line_spacing = 1.35

        p.paragraph_format.space_after = Pt(6)

        clean_inline = line_clean.replace(
            "$",
            ""
        )

        # INLINE LATEX CLEANING

        clean_inline = clean_inline.replace(
            r"\beta",
            "β"
        )

        clean_inline = clean_inline.replace(
            r"\alpha",
            "α"
        )

        clean_inline = clean_inline.replace(
            r"\sigma",
            "σ"
        )

        clean_inline = clean_inline.replace(
            r"\theta",
            "θ"
        )

        clean_inline = clean_inline.replace(
            r"\lambda",
            "λ"
        )

        clean_inline = clean_inline.replace(
            r"\omega",
            "ω"
        )

        clean_inline = clean_inline.replace(
            r"\gamma",
            "γ"
        )

        clean_inline = clean_inline.replace(
            r"\delta",
            "δ"
        )

        clean_inline = clean_inline.replace(
            r"\times",
            "×"
        )

        clean_inline = clean_inline.replace(
            r"\%",
            "%"
        )

        clean_inline = clean_inline.replace(
            r"\geq",
            "≥"
        )

        clean_inline = clean_inline.replace(
            r"\leq",
            "≤"
        )

        clean_inline = clean_inline.replace(
            r"\neq",
            "≠"
        )

        clean_inline = clean_inline.replace(
            r"\approx",
            "≈"
        )

        # SIMPLE FRACTION CLEANING

        clean_inline = re.sub(
            r'\\frac\{([^{}]+)\}\{([^{}]+)\}',
            r'(\1/\2)',
            clean_inline
        )

        # INLINE SUBSCRIPT FIXES

        clean_inline = clean_inline.replace(
            "_e",
            "ₑ"
        )

        clean_inline = clean_inline.replace(
            "_f",
            "բ"
        )

        clean_inline = clean_inline.replace(
            "_m",
            "ₘ"
        )

        clean_inline = clean_inline.replace(
            "_0",
            "₀"
        )

        clean_inline = clean_inline.replace(
            "_1",
            "₁"
        )

        clean_inline = clean_inline.replace(
            "_2",
            "₂"
        )

        clean_inline = clean_inline.replace(
            "_3",
            "₃"
        )

        clean_inline = clean_inline.replace(
            "_t",
            "ₜ"
        )

        # INLINE SUPERSCRIPT FIXES

        clean_inline = clean_inline.replace(
            "^2",
            "²"
        )

        clean_inline = clean_inline.replace(
            "^3",
            "³"
        )

        clean_inline = clean_inline.replace(
            "^n",
            "ⁿ"
        )

        clean_inline = clean_inline.replace(
            "^t",
            "ᵗ"
        )

        # REMOVE EXTRA LATEX WRAPPERS

        clean_inline = clean_inline.replace(
            "{",
            ""
        )

        clean_inline = clean_inline.replace(
            "}",
            ""
        )

        clean_inline = clean_inline.replace(
            "\\",
            ""
        )

        clean_inline = clean_inline.replace(
            "^",
            ""
        )

        clean_inline = clean_inline.replace(
            "_",
            ""
        )

        run = p.add_run(
            clean_inline
        )

        run.font.size = Pt(12)

        run.font.name = (
            "Times New Roman"
        )

        i += 1

    # SAVE FILE

    filename = (
        f"report_{uuid.uuid4()}.docx"
    )

    filepath = os.path.join(
        os.getcwd(),
        filename
    )

    doc.save(filepath)

    background_tasks.add_task(
        os.remove,
        filepath
    )

    return FileResponse(

        filepath,

        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),

        filename="MSPK_Report.docx"
    )

# =========================================
# HEALTH
# =========================================

@app.get("/health")

def health():

    return {
        "status": "ok"
    }

# =========================================
# HOME
# =========================================

@app.get("/")

def home():

    return {

        "status": "online",

        "message": "MSPK BACKEND RUNNING"

    }

# =========================================
# MAIN
# =========================================

if __name__ == "__main__":

    import uvicorn

    logger.info("=" * 60)

    logger.info(
        "MSPK AI to Word Converter – Backend v5.4"
    )

    logger.info("=" * 60)

    uvicorn.run(
        app,
        host="0.0.0.0",
        port=8000
    )